#!/usr/bin/env python3
"""
update_docx.py — Inject test results and screenshots back into the original Word document.

Appends a "Test Execution Summary" section at the end of the original .docx,
adds PASS/FAIL status (colour-coded) next to each test case, and embeds
representative screenshots.

Usage:
    python scripts/update_docx.py docs/input.docx reports/results.json
    python scripts/update_docx.py docs/input.docx reports/results.json --output reports/report_input.docx
    python scripts/update_docx.py docs/input.docx reports/results.json \
        --output reports/report_input.docx \
        --screenshots-dir ./reports/screenshots
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------

GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED   = RGBColor(0xDC, 0x26, 0x26)
GREY  = RGBColor(0x9C, 0xA3, 0xAF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_GREEN_BG = "DCF2E4"   # For table cell shading
LIGHT_RED_BG   = "FEE2E2"
LIGHT_GREY_BG  = "F3F4F6"


# ---------------------------------------------------------------------------
# 字體保留輔助函式（解決回寫時字型跑掉的問題）
# ---------------------------------------------------------------------------

def _capture_font_style(cell) -> dict:
    """擷取儲存格中第一個 run 的字體資訊，作為後續回寫的基準。"""
    style_info: dict = {"name": None, "size": None, "east_asia": None}
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                style_info["name"] = run.font.name
            if run.font.size:
                style_info["size"] = run.font.size
            # 嘗試讀取東亞字體
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                ea = rFonts.get(qn("w:eastAsia"))
                if ea:
                    style_info["east_asia"] = ea
            if style_info["name"] or style_info["east_asia"]:
                return style_info
    return style_info


def _apply_font_style(
    run, style_info: dict,
    color: RGBColor | None = None,
    bold: bool = False,
) -> None:
    """將記錄的字體樣式套用到新 run 上，確保字型一致。"""
    if style_info.get("name"):
        run.font.name = style_info["name"]
    if style_info.get("size"):
        run.font.size = style_info["size"]
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    # 設定東亞字體（中文必要）
    ea_name = style_info.get("east_asia") or style_info.get("name")
    if ea_name:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), ea_name)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex: str) -> None:
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_coloured_heading(doc: Document, text: str, level: int = 1, colour: RGBColor = DARK) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = colour


def status_colour(status: str) -> tuple[RGBColor, str]:
    """Return (text_colour, bg_hex) for a status string."""
    s = status.lower()
    if s == "passed":
        return GREEN, LIGHT_GREEN_BG
    if s == "failed":
        return RED, LIGHT_RED_BG
    return GREY, LIGHT_GREY_BG


def status_label(status: str) -> str:
    return {"passed": "✅ PASS", "failed": "❌ FAIL", "skipped": "⏭ SKIP"}.get(status.lower(), status.upper())


# ---------------------------------------------------------------------------
# In-place table updater (preserves original Word formatting)
# ---------------------------------------------------------------------------

def _unique_cells(row) -> list:
    """Return de-duplicated cells for a row (handles merged cells)."""
    cells = []
    seen: set[int] = set()
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell)
    return cells


def _update_result_cell(cell, status: str) -> None:
    """更新 □通過 / □失敗 checkboxes，保留原始字體樣式。"""
    original_font = _capture_font_style(cell)
    for paragraph in cell.paragraphs:
        full_text = paragraph.text
        if "通過" in full_text:
            for run in paragraph.runs:
                run.text = ""
            if status == "passed":
                r = paragraph.add_run("☑通過")
                _apply_font_style(r, original_font, color=GREEN, bold=True)
            else:
                r = paragraph.add_run("□通過")
                _apply_font_style(r, original_font, color=GREY)
        elif "失敗" in full_text:
            for run in paragraph.runs:
                run.text = ""
            if status == "failed":
                r = paragraph.add_run("☑失敗")
                _apply_font_style(r, original_font, color=RED, bold=True)
            else:
                r = paragraph.add_run("□失敗")
                _apply_font_style(r, original_font, color=GREY)
    # Fallback: no checkbox text found — append a status label
    has_checkbox = any("通過" in p.text or "失敗" in p.text for p in cell.paragraphs)
    if not has_checkbox:
        label = status_label(status)
        colour, _ = status_colour(status)
        p = cell.add_paragraph()
        r = p.add_run(label)
        _apply_font_style(r, original_font, color=colour, bold=True)


def _fill_record_cell(cell, results: dict) -> None:
    """Write a brief summary into the 測試紀錄 cell."""
    summary = results.get("summary", {})
    started = summary.get("started_at", "")[:19].replace("T", " ")
    passed = summary.get("passed", 0)
    failed = summary.get("failed", 0)
    total = summary.get("total", 0)
    duration = summary.get("duration_seconds", 0)
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    for run in p.runs:
        run.text = ""
    r = p.add_run(
        f"自動測試 | {started} | "
        f"{passed}/{total} 通過 | 耗時 {duration}s"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN if failed == 0 else RED


def update_tables_in_place(
    doc: "Document",
    results: dict,
    screenshot_dir: Path,
) -> int:
    """
    Locate test-case tables and update the 測試結果 column in-place
    (□ → ☑) while preserving original formatting.
    Returns the number of rows updated.
    """
    test_cases: list[dict] = results.get("test_cases", [])
    result_by_name: dict[str, dict] = {tc.get("name", ""): tc for tc in test_cases}

    updated = 0
    tc_global_idx = 0

    for table in doc.tables:
        # ── Detect column layout from header row ─────────────────────
        result_col = -1
        name_col = -1
        for row in table.rows:
            unique = _unique_cells(row)
            cell_texts = [c.text.strip().replace("**", "") for c in unique]
            for i, ct in enumerate(cell_texts):
                if "測試結果" in ct:
                    result_col = i
                if "測試項目" in ct:
                    name_col = i
            if result_col >= 0:
                break
        if result_col < 0:
            continue

        # ── Update data rows ───────────────────────────────────────
        # Collect header keywords to skip header/meta rows
        _skip_kw = re.compile(r"個案編號|項目序|測試目標|測試紀錄|個案名稱")

        for row in table.rows:
            unique = _unique_cells(row)
            if len(unique) <= result_col:
                continue
            first_text = unique[0].text.strip().replace("**", "")

            # Handle 測試紀錄 row
            if "測試紀錄" in first_text and len(unique) >= 2:
                _fill_record_cell(unique[1], results)
                continue

            # Skip header / meta rows
            if _skip_kw.search(first_text):
                continue

            # Identify data rows: the result cell must contain checkbox markers
            result_cell_text = unique[result_col].text if result_col < len(unique) else ""
            has_checkbox = "通過" in result_cell_text or "失敗" in result_cell_text
            tc_name = ""
            if 0 <= name_col < len(unique):
                tc_name = unique[name_col].text.strip()

            if not has_checkbox or not tc_name:
                continue

            # Match test result by name or global index
            tc_result = result_by_name.get(tc_name)
            if not tc_result and tc_global_idx < len(test_cases):
                tc_result = test_cases[tc_global_idx]
            tc_global_idx += 1
            if not tc_result:
                continue

            _update_result_cell(unique[result_col], tc_result.get("status", "skipped"))
            updated += 1

    return updated


# ---------------------------------------------------------------------------
# Summary section builder
# ---------------------------------------------------------------------------

def append_summary_section(
    doc: Document,
    results: dict,
    screenshot_dir: Path,
) -> None:
    meta = results.get("meta", {})
    summary = results.get("summary", {})
    test_cases: list[dict] = results.get("test_cases", [])

    total = summary.get("total", len(test_cases))
    passed = summary.get("passed", 0)
    failed = summary.get("failed", 0)
    skipped = summary.get("skipped", 0)
    duration = summary.get("duration_seconds", 0)
    started_at = summary.get("started_at", "")
    pass_rate = round(passed / total * 100) if total else 0

    # Page break before the new section
    doc.add_page_break()

    # ── Section heading ────────────────────────────────────────────────────
    add_coloured_heading(doc, "測試執行報告 (Test Execution Summary)", level=1, colour=DARK)
    doc.add_paragraph(
        f"自動產生 | 執行時間：{started_at[:19].replace('T', ' ')} UTC  |  "
        f"通過率：{pass_rate}%  |  共 {total} 案例  |  耗時 {duration}s"
    ).style.font.size = Pt(9)

    add_horizontal_rule(doc)

    # ── Stats overview table ───────────────────────────────────────────────
    doc.add_heading("執行統計", level=2)
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    headers = ["總案例", "通過", "失敗", "略過", "通過率"]
    values = [str(total), str(passed), str(failed), str(skipped), f"{pass_rate}%"]
    colours = [DARK, GREEN, RED, GREY, GREEN if pass_rate >= 80 else RED]
    for i, (h, v, c) in enumerate(zip(headers, values, colours)):
        hdr_cell = tbl.rows[0].cells[i]
        hdr_cell.text = h
        hdr_cell.paragraphs[0].runs[0].font.bold = True
        val_cell = tbl.rows[1].cells[i]
        run = val_cell.paragraphs[0].add_run(v)
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = c
    doc.add_paragraph()  # spacing

    # ── Fail list (if any) ────────────────────────────────────────────────
    failed_cases = [tc for tc in test_cases if tc.get("status") == "failed"]
    if failed_cases:
        p = doc.add_paragraph()
        p.add_run("❌ 失敗案例清單").bold = True
        for tc in failed_cases:
            try:
                item = doc.add_paragraph(style="List Bullet")
            except KeyError:
                item = doc.add_paragraph()
                item.paragraph_format.left_indent = Pt(18)
            r = item.add_run(f"• {tc['id']} — {tc['name']}")
            r.font.color.rgb = RED

    add_horizontal_rule(doc)

    # ── Per-test-case results ──────────────────────────────────────────────
    doc.add_heading("逐案測試結果", level=2)

    for tc in test_cases:
        tc_id = tc.get("id", "")
        tc_name = tc.get("name", "")
        tc_status = tc.get("status", "skipped")
        tc_dur = tc.get("duration_seconds", 0)
        tc_error = tc.get("error") or ""
        steps: list[dict] = tc.get("steps", [])

        text_col, bg_hex = status_colour(tc_status)
        label = status_label(tc_status)

        # Case heading row
        p = doc.add_paragraph()
        r = p.add_run(f"  {label}  ")
        r.font.bold = True
        r.font.color.rgb = text_col
        r = p.add_run(f"  {tc_id}: {tc_name}  ({tc_dur}s)")
        r.font.bold = True
        r.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(12)

        # Error message
        if tc_error:
            ep = doc.add_paragraph()
            er = ep.add_run("  ⚠ " + tc_error[:500])
            er.font.color.rgb = RED
            er.font.size = Pt(9)
            er.font.italic = True

        # Steps table
        if steps:
            st = doc.add_table(rows=1, cols=4)
            st.style = "Table Grid"
            hdr_labels = ["#", "動作", "對象 / 數值", "結果"]
            for i, hl in enumerate(hdr_labels):
                cell = st.rows[0].cells[i]
                cell.text = hl
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_shading(cell, "F9FAFB")

            for idx, step in enumerate(steps):
                row = st.add_row()
                sc, sbg = status_colour(step.get("status", "passed"))
                step_label = "✓" if step.get("status") == "passed" else "✗"
                cells = row.cells
                cells[0].text = str(idx + 1)
                cells[1].text = step.get("action", "")
                target = str(step.get("target", ""))
                val = str(step.get("value", ""))
                cells[2].text = (target + (" → " + val if val else ""))[:120]
                r2 = cells[3].paragraphs[0].add_run(step_label)
                r2.font.color.rgb = sc
                r2.font.bold = True
                if step.get("status") == "failed":
                    for c in cells:
                        set_cell_shading(c, LIGHT_RED_BG)

        # 截圖嵌入（不限數量，去除重複，加上描述）
        seen_screenshots: set[str] = set()
        for step in steps:
            ss_name = step.get("screenshot")
            if not ss_name or ss_name in seen_screenshots:
                continue
            seen_screenshots.add(ss_name)
            img_path = screenshot_dir / ss_name
            if not img_path.exists():
                continue
            try:
                doc.add_paragraph()
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(str(img_path), width=Inches(5.5))
                # 截圖描述：動作 + 對象，50 字內
                action_desc = step.get("action", "")
                target_desc = str(step.get("target", ""))[:40]
                caption_text = f"  截圖: [{action_desc}] {target_desc}"
                cap = doc.add_paragraph(caption_text[:60])
                for r in cap.runs:
                    r.font.size = Pt(8)
                    r.font.italic = True
            except Exception as e:
                print(f"  [WARN] Could not embed screenshot {ss_name}: {e}", file=sys.stderr)

        add_horizontal_rule(doc)

    # ── Footer note ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"本報告由 browser-test-reporter skill 自動產生。"
        f"完整截圖與錄影請參閱隨附的 report.html。\n"
        f"Base URL: {meta.get('base_url', '—')}  |  環境: {meta.get('environment', '—')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = GREY
    footer_run.font.italic = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Inject test results into the original Word document",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx", help="Path to the original input .docx file")
    parser.add_argument("results", help="Path to results.json (from run_tests.py)")
    parser.add_argument("--output", "-o", default="", help="Output .docx path (default: reports/report_<input>.docx)")
    parser.add_argument("--screenshots-dir", default="", help="Directory containing screenshot files")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    results_path = Path(args.results)

    for p in (docx_path, results_path):
        if not p.exists():
            print(f"[ERROR] File not found: {p}", file=sys.stderr)
            sys.exit(1)

    out_path = Path(args.output) if args.output else results_path.parent / f"report_{docx_path.name}"

    # Resolve screenshot dir
    if args.screenshots_dir:
        screenshot_dir = Path(args.screenshots_dir)
    else:
        screenshot_dir = results_path.parent / "screenshots"
    if not screenshot_dir.exists():
        print(f"[WARN] Screenshots directory not found: {screenshot_dir}", file=sys.stderr)
        screenshot_dir = Path(".")

    # Load results
    results = json.loads(results_path.read_text(encoding="utf-8"))

    # Copy original doc so we don't destructively modify it
    shutil.copy2(docx_path, out_path)
    print(f"[1/2] Copied {docx_path.name} → {out_path.name}")

    # Open and modify
    doc = Document(str(out_path))

    # Step 1: Update test-result tables in-place (□ → ☑)
    print("[2/3] Updating test result tables in-place...")
    n_updated = update_tables_in_place(doc, results, screenshot_dir)
    print(f"      Updated {n_updated} table cell(s)")

    # Step 2: Append detailed summary section with screenshots
    print("[3/3] Appending test execution summary...")
    append_summary_section(doc, results, screenshot_dir)
    doc.save(str(out_path))

    total = results.get("summary", {}).get("total", 0)
    passed = results.get("summary", {}).get("passed", 0)
    failed = results.get("summary", {}).get("failed", 0)

    print(f"\n✅ Updated Word document → {out_path.resolve()}")
    print(f"   {total} cases: {passed} passed / {failed} failed")


if __name__ == "__main__":
    main()
